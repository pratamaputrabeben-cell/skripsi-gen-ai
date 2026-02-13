import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import random
import re

# --- 1. KONFIGURASI ENGINE ---
ALL_KEYS = st.secrets.get("GEMINI_API_KEYS", [st.secrets.get("GEMINI_API_KEY", "")])
if 'db' not in st.session_state: st.session_state['db'] = {}
if 'user_data' not in st.session_state:
    st.session_state['user_data'] = {"topik": "", "lokasi": "Puskesmas Bandar Jaya", "kota": "Lahat", "nama": ""}

def inisialisasi_ai():
    key_aktif = random.choice(ALL_KEYS)
    genai.configure(api_key=key_aktif)
    try:
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        for target in ['gemini-1.5-flash', 'gemini-1.5-pro']:
            for real_model in available_models:
                if target in real_model: return genai.GenerativeModel(real_model)
        return genai.GenerativeModel(available_models[0])
    except: return genai.GenerativeModel('gemini-1.5-flash')

# --- 2. FUNGSI MEMBERSIHKAN & MENGURUTKAN TEKS ---
def bersihkan_dan_urutkan(teks):
    teks = re.sub(r"^(Tentu|Berikut|Ini adalah|Sesuai).*?\n", "", teks, flags=re.IGNORECASE)
    teks = teks.replace("&nbsp;", " ").replace("**", "").replace("---", "")
    if "DAFTAR PUSTAKA" in teks.upper():
        parts = re.split(r"DAFTAR PUSTAKA", teks, flags=re.IGNORECASE)
        konten_utama = parts[0]
        pustaka_raw = parts[1].strip().split('\n')
        pustaka_clean = sorted([p.strip() for p in pustaka_raw if len(p.strip()) > 10])
        return konten_utama.strip(), pustaka_clean
    return teks.strip(), []

# --- 3. FUNGSI RAPIKAN WORD (STANDAR AKADEMIK + AUTO TABLE) ---
def buat_dokumen_rapi(judul_bab, isi_teks):
    doc = Document()
    # Set Margin 4-3-3-3 
    for sec in doc.sections:
        sec.left_margin, sec.top_margin = Cm(4), Cm(3)
        sec.right_margin, sec.bottom_margin = Cm(3), Cm(3)

    style = doc.styles['Normal']
    style.font.name, style.font.size = 'Times New Roman', Pt(12)
    
    konten, daftar_pustaka = bersihkan_dan_urutkan(isi_teks)
    lines = konten.split('\n')

    # Header Bab [cite: 312, 1124]
    head = doc.add_heading(judul_bab.upper(), 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in head.runs:
        run.font.name, run.font.size, run.bold, run.font.color.rgb = 'Times New Roman', Pt(14), True, None
    
    # LOGIKA AUTO-TABLE: Deteksi jika ada tabel kuesioner [cite: 1729, 1826]
    if "|" in konten and "No" in konten:
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['No', 'Pernyataan', 'SS', 'S', 'N', 'TS', 'STS']
        for i, h_text in enumerate(headers):
            hdr_cells[i].text = h_text
        
        for line in lines:
            if "|" in line and "No" not in line and "---" not in line:
                parts = [p.strip() for p in line.split('|') if p.strip() != ""]
                if len(parts) >= 2:
                    row_cells = table.add_row().cells
                    row_cells[0].text = parts[0] # Kolom No
                    row_cells[1].text = parts[1] # Kolom Pernyataan
                    # Isi kolom jawaban SS-STS jika ada centang [cite: 1691, 1695]
                    for i in range(2, min(len(parts), 7)):
                        row_cells[i].text = parts[i]
    else:
        # Proses Konten Utama (Paragraf Biasa)
        for p_text in lines:
            t = p_text.strip()
            if t:
                p = doc.add_paragraph()
                fmt = p.paragraph_format
                fmt.line_spacing, fmt.alignment = 1.5, WD_ALIGN_PARAGRAPH.JUSTIFY
                match_num = re.match(r"^(\d+\.\d+(\.\d+)?)\s*(.*)", t)
                if match_num:
                    run = p.add_run(t)
                    run.bold = True
                    level = match_num.group(1).count('.')
                    fmt.left_indent = Inches(0.2 * level)
                else:
                    p.add_run(t)
                    fmt.first_line_indent = Inches(0.5)

    # Proses Daftar Pustaka [cite: 3, 290]
    if daftar_pustaka:
        doc.add_page_break()
        dp_head = doc.add_heading("DAFTAR PUSTAKA", 0)
        dp_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in dp_head.runs:
            run.font.name, run.font.size, run.bold, run.font.color.rgb = 'Times New Roman', Pt(14), True, None
        for ref in daftar_pustaka:
            p = doc.add_paragraph(ref)
            fmt = p.paragraph_format
            fmt.line_spacing, fmt.left_indent, fmt.first_line_indent = 1.5, Inches(0.5), Inches(-0.5)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 4. UI STREAMLIT ---
st.set_page_config(page_title="SkripsiGen Pro v8.51", layout="wide", page_icon="🎓")

with st.sidebar:
    st.header("🛡️ Pusat Kalibrasi")
    st.success("✅ Auto-Table Mode: ACTIVE")
    st.success("✅ Margin 4333: ACTIVE")
    
    nama_user = st.text_input("👤 Nama Mahasiswa:", value=st.session_state['user_data']['nama'])
    st.session_state['user_data']['nama'] = nama_user
    user_lic = st.text_input("🔑 Kode Lisensi PRO:", type="password")
    
    def gen_lic(n):
        d = datetime.now().strftime("%d%m")
        nm = n.split(' ')[0].upper() if n else "USER"
        return f"PRO-{nm}-{d}-SKR"

    st.divider()
    with st.expander("🛠️ OWNER PANEL"):
        pw = st.text_input("Admin Password:", type="password")
        if pw == "RAHASIA-BEBEN-2026": 
            pbl = st.text_input("Nama Pembeli:")
            if st.button("Generate License ✨"): st.code(gen_lic(pbl))
    
    if st.button("🗑️ Reset Semua"):
        st.session_state['db'] = {}
        st.rerun()

st.title("🎓 SkripsiGen Pro v8.51")
st.caption("Standard: Professional Academic Formatting | Anti-Plagiarism Engine")

col1, col2 = st.columns(2)
with col1:
    topik = st.text_input("📝 Judul Skripsi:", value=st.session_state['user_data']['topik'])
    st.session_state['user_data']['topik'] = topik
    lokasi = st.text_input("📍 Lokasi:", value=st.session_state['user_data']['lokasi'])
    st.session_state['user_data']['lokasi'] = lokasi
with col2:
    kota = st.text_input("🏙️ Kota:", value=st.session_state['user_data']['kota'])
    st.session_state['user_data']['kota'] = kota
    metode = st.selectbox("🔬 Metode:", ["Kuantitatif", "Kualitatif", "R&D"])

st.divider()
pil_bab = st.selectbox("📄 Pilih Bagian:", ["Bab 1", "Bab 2", "Bab 3", "Bab 4", "Bab 5", "Lampiran: Surat Izin, Kuesioner & Kisi-kisi"])

def jalankan_proses(target_bab=None, catatan_dosen=""):
    bab_aktif = target_bab if target_bab else pil_bab
    if topik and nama_user:
        with st.spinner(f"Menyusun & Kalibrasi {bab_aktif}..."):
            try:
                model = inisialisasi_ai()
                # GLOBAL MANTRA: Memaksa Tabel Markdown [cite: 1729, 1826]
                inst = """
                Gunakan standar akademik profesional. 
                WAJIB: Sajikan data klasifikasi, instrumen penelitian, 
                definisi operasional, atau data statistik dalam **Markdown Table**. 
                Khusus kuesioner, gunakan kolom: No, Pernyataan, SS, S, N, TS, STS.
                Referensi RIIL 2018-2026, APA 7th.
                """
                prompt = f"{inst}\nSusun {bab_aktif} skripsi {metode} judul '{topik}' di {lokasi}, {kota}. Revisi: {catatan_dosen}."
                res = model.generate_content(prompt)
                st.session_state['db'][bab_aktif] = res.text
                st.success(f"{bab_aktif} Berhasil Disusun!")
                st.rerun()
            except: st.error("Server sibuk, klik sekali lagi!")
    else: st.warning("Nama dan Judul wajib diisi!")

if st.button("🚀 Susun & Kalibrasi Sekarang"):
    jalankan_proses()

if st.session_state['db']:
    st.divider()
    for b in sorted(st.session_state['db'].keys()):
        content = st.session_state['db'][b]
        with st.container(border=True):
            st.markdown(f"### 📄 {b}")
            is_trial, is_pro = b in ["Bab 1", "Bab 2"], user_lic == gen_lic(nama_user)
            if "Lampiran" not in b:
                cat = st.text_area(f"✍️ Revisi Dosen {b}:", key=f"rev_{b}")
                if st.button(f"🔄 Jalankan Revisi {b}", key=f"br_{b}"): jalankan_proses(target_bab=b, catatan_dosen=cat)
            with st.expander(f"Buka Draf {b}"):
                st.text_area("Pratinjau Teks:", value=content, height=300, disabled=True)
                if is_trial or is_pro:
                    data_word = buat_dokumen_rapi(b, content)
                    st.download_button(f"📥 Download {b} (Word Rapi)", data=data_word, file_name=f"{b}.docx", key=f"dl_{b}")
                else:
                    st.error(f"🔑 {b} Terkunci (Mode PRO)")
                    st.link_button("💬 Hubungi Admin (Aktivasi)", f"https://wa.me/6281273347072?text=Mau%20Lisensi%20PRO%20{nama_user}")
