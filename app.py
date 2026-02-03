import streamlit as st
import requests
import google.generativeai as genai
from docx import Document
from io import BytesIO

# --- KONFIGURASI API AMAN ---
# Pastikan Anda sudah mengisi GEMINI_API_KEY di menu Secrets Streamlit
try:
    api_key = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=api_key)
    # Langsung pakai model spesifik (Tanpa fungsi get_active_model yang error tadi)
    model = genai.GenerativeModel('gemini-1.5-flash')
except Exception as e:
    st.error(f"Masalah pada API Key atau Konfigurasi: {e}")
    st.stop()

# --- FUNGSI PENDUKUNG ---
def search_real_papers(query):
    # Mencari jurnal asli dari Semantic Scholar
    url = f"https://api.semanticscholar.org/graph/v1/paper/search?query={query}&limit=3&fields=title,authors,year,abstract,url"
    try:
        response = requests.get(url, timeout=5)
        return response.json().get('data', [])
    except:
        return []

def create_docx(judul, konten):
    doc = Document()
    doc.add_heading(judul, 0)
    doc.add_paragraph(konten)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- TAMPILAN WEBSITE ---
st.set_page_config(page_title="Penyusun Skripsi Otomatis", layout="wide")
st.title("🎓 SkripsiGen Pro")
st.info("Aplikasi sudah aktif. Masukkan topik skripsi Anda di bawah ini.")

topik = st.text_input("Masukkan Topik/Judul Skripsi:", placeholder="Contoh: Pengaruh gadget pada mahasiswa")
bahasa = st.selectbox("Pilih Bahasa:", ["Indonesia", "English"])

if st.button("Generate Draf Bab 1 ✨"):
    if topik:
        with st.spinner("Sedang menyusun draf menggunakan AI..."):
            # 1. Cari Jurnal
            papers = search_real_papers(topik)
            context = ""
            if papers:
                for p in papers:
                    context += f"Judul: {p['title']}\nAbstrak: {p['abstract']}\nURL: {p['url']}\n\n"
            
            # 2. Buat Prompt untuk AI
            prompt = f"Buatkan draf Bab 1 Skripsi formal dalam bahasa {bahasa} dengan judul '{topik}'. Gunakan referensi riil ini jika tersedia:\n{context}. Jika tidak ada, gunakan teori akademik yang umum. Sertakan Latar Belakang dan Daftar Pustaka."
            
            try:
                # 3. Panggil AI
                response = model.generate_content(prompt)
                hasil_teks = response.text
                
                # 4. Tampilkan Hasil
                col1, col2 = st.columns(2)
                with col1:
                    st.subheader("📝 Draf Bab 1")
                    st.write(hasil_teks)
                    file_word = create_docx(topik, hasil_teks)
                    st.download_button("📥 Download File Word", data=file_word, file_name=f"Draf_Skripsi_{topik}.docx")
                
                with col2:
                    st.subheader("📚 Referensi Jurnal Riil")
                    if papers:
                        for p in papers:
                            st.info(f"**{p['title']}** ({p['year']})\n[Link Jurnal]({p['url']})")
                    else:
                        st.warning("Jurnal spesifik tidak ditemukan di database. Draf dibuat menggunakan basis data pengetahuan AI.")
            
            except Exception as e:
                st.error(f"Gagal generate konten: {e}")
    else:
        st.warning("Tuliskan judulnya dulu ya!")
